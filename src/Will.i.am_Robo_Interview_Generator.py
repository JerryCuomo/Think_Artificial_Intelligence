# Source: "Think Artificial Intelligence" by Jerry Cuomo, 2024
# Purpose: Educational code examples from the book.
# Copyright © 2024 Jerry Cuomo. All rights reserved.
#
# This code uses OpenAI's GPT-4 to automatically generate responses to interview questions in a specified voice, processes a given CSV dataset of questions and answers rated with a '1' (indicating high relevance or importance), and compiles the responses along with the original answers into a Word document. The script demonstrates the integration of AI in automating content creation, showcasing how AI can be used to simulate interview responses and document compilation.
#
# About: The script exemplifies the application of AI for educational and preparatory purposes, specifically in the context of job interviews. It leverages the OpenAI GPT-4 model to generate articulate and concise answers to interview questions, reflecting the specified persona of a robot named Robo. This process not only highlights the capabilities of AI in understanding and generating human-like responses but also showcases its potential in educational tools and resources. Additionally, the script demonstrates practical skills in Python programming, including working with CSV files, handling document creation, and utilizing environmental variables for API keys.
#
# Setup: Requires Python with libraries: langchain_openai for interacting with OpenAI's GPT-4, langchain_core for prompt management, csv for reading CSV files, and python-docx for Word document manipulation. Install dependencies using 'pip install langchain_openai langchain_core python-docx'.

import os
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
import csv
from docx import Document

# Constants
OPENAI_API_KEY =  "INSERT_YOUR_OPENAI_KEY_HERE"
DATASET_FILE_PATH = "../datasets/will.i.am-interview-question-answers.csv"
WORD_DOCUMENT_PATH = "../datasets/mock_interview.docx"
LANGUAGE = "an articulate humanoid robot voice, identifying as a girl named Robo"
NUM_QUESTIONS = 10

# Set environment variable for OpenAI API key
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY

# Initialize the language model and document
llm = ChatOpenAI(temperature=0)
document = Document()

# Define a prompt template
prompt_template = ChatPromptTemplate.from_messages([
    ("system", "Please respond as a question in {language} in 25 words or less. {validate}"),
    ("human", "{input}")
])

# Function to process a single question
def process_question(input_question, validate=""):
    response = prompt_template | llm
    processed_response = response.invoke({"input": input_question, "validate": validate, "language": LANGUAGE})
    print(processed_response.content)
    return processed_response.content

# Read dataset and process questions
with open(DATASET_FILE_PATH, mode='r', encoding='ISO-8859-1') as file:
    csv_reader = csv.DictReader(file)
    questions_processed = 0
    endnotes = []

    for row in csv_reader:
        if row['Rating'] == '1':
            if questions_processed == 0:
                # Process the first question for introduction
                intro_response = process_question(row['Question'])
                document.add_paragraph(f"Introduction: {intro_response}\n")
            else:
                # Process subsequent questions
                processed_question = process_question(row['Question'])
                document.add_paragraph(f"Q{questions_processed}: {processed_question}")
                document.add_paragraph(f"A{questions_processed}: {row['Answer']}")
                endnotes.append(row['Reference'])

            questions_processed += 1
            if questions_processed > NUM_QUESTIONS:
                break

# Add endnotes
document.add_page_break()
document.add_paragraph("Endnotes:")
for i, note in enumerate(endnotes, start=1):
    document.add_paragraph(f"{i}. {note}")

# Save the document
document.save(WORD_DOCUMENT_PATH)

print(f"Mock interview document with {NUM_QUESTIONS} questions saved as '{WORD_DOCUMENT_PATH}'")
